import streamlit as st
from langchain.prompts import PromptTemplate
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.output_parsers import PydanticOutputParser
from pydantic import BaseModel, Field
from PyPDF2 import PdfReader
import docx
from langchain.schema import Document
from langchain_core.callbacks.base import BaseCallbackHandler
import io
from bson.objectid import ObjectId
from PIL import Image
import base64
from urllib.parse import urlencode

from agent import get_agent_executor, get_search_tool
from ext_tools.qa_tool import qa_generation
from ext_tools.instant_rag import create_rag_tool

from utils.database import (
    get_chat_sessions,
    create_chat_session,
    prepare_chat_history,
    add_message_to_session,
    update_session_name
)

def inject_custom_css():
    st.markdown(
        """
        <style>
        .chat-list-item {
            padding: 6px 10px;
            margin: 2px 0;
            border-radius: 5px;
            cursor: pointer;
            transition: background-color 0.2s;
            border: 1px solid rgba(49, 51, 63, 0.15);
            text-align: left;
            display: flex;
            align-items: center;
            width: 100%;
            background-color: transparent;
        }

        .chat-list-item:hover {
            background-color: rgba(151, 166, 195, 0.12);
        }

        .chat-list-item.active {
            background-color: rgba(151, 166, 195, 0.25);
            border-color: rgb(49, 51, 63);
        }

        .chat-title {
            font-size: 14px;
            white-space: nowrap;
            overflow: hidden;
            text-overflow: ellipsis;
            color: rgb(33, 37, 41);  /* Darker text for readability */
        }
        </style>
        """,
        unsafe_allow_html=True
    )


class LambdaStreamlitLoader:
    def __init__(self, uploaded_file) -> None:
        self.uploaded_file = uploaded_file
        self.file_name = uploaded_file.name
        if "." in self.file_name:
            *_, self.ext = self.file_name.rsplit(".", 1)
            self.ext = self.ext.lower()
        else:
            self.ext = ""

    def process_image(self):
        """Process image files and convert to base64 for storage/display."""
        try:
            image = Image.open(io.BytesIO(self.uploaded_file.getvalue()))
            buffered = io.BytesIO()
            image.save(buffered, format=image.format)
            img_str = base64.b64encode(buffered.getvalue()).decode()
            return Document(
                page_content=f"![{self.file_name}](data:image/{self.ext};base64,{img_str})",
                metadata={"source": self.file_name, "type": "image"}
            )
        except Exception as e:
            st.error(f"Error processing image '{self.file_name}': {e}")
            return None

    def lazy_load(self):
        """Yields Document objects from the uploaded file."""
        try:
            file_content = self.uploaded_file.getvalue()
            if not file_content:
                st.warning(f"File '{self.file_name}' appears to be empty.")
                return

            if self.ext in ["jpg", "jpeg", "png", "gif"]:
                img_doc = self.process_image()
                if img_doc:
                    yield img_doc

            elif self.ext in ["docx", "doc"]:
                doc = docx.Document(io.BytesIO(file_content))
                for paragraph in doc.paragraphs:
                    lines = paragraph.text.split("\n")
                    for line in lines:
                        line = line.strip()
                        if line:
                            yield Document(page_content=line, metadata={"source": self.file_name})

            elif self.ext == "pdf":
                reader = PdfReader(io.BytesIO(file_content))
                if not reader.pages:
                    st.warning(f"Could not read any pages from PDF '{self.file_name}'.")
                    return
                for i, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text:
                        lines = text.split("\n")
                        for line in lines:
                            line = line.strip()
                            if line:
                                yield Document(
                                    page_content=line,
                                    metadata={"source": self.file_name, "page": i + 1}
                                )
            else:
                st.warning(f"Unsupported file format: '{self.ext if self.ext else 'Unknown'}'")
                return

        except Exception as e:
            st.error(f"Error processing file '{self.file_name}': {e}")
            return

def get_base_title(unique_title: str) -> str:
    """Extracts the base title from a unique title (title_sessionid)."""
    if not unique_title or not isinstance(unique_title, str):
        return "Chat"
    parts = unique_title.rsplit('_', 1)
    if len(parts) == 2 and len(parts[1]) == 24:
        try:
            ObjectId(parts[1])
            return parts[0]
        except Exception:
            return unique_title
    else:
        return unique_title

class TitleParser(BaseModel):
    title: str = Field(description="Title of the chat session")

title_parser = PydanticOutputParser(pydantic_object=TitleParser)

def generate_title_llm(first_message: str) -> str:
    """Generates the base title (without unique ID)."""
    prompt = PromptTemplate(
        template="Based on the given message, suggest a suitable title for the chat (max 5 words).\nMessage: {message}\n{format_instructions}",
        input_variables=["message"],
        partial_variables={"format_instructions": title_parser.get_format_instructions()}
    )
    llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.2)
    chain = prompt | llm | title_parser
    try:
        result = chain.invoke({"message": first_message})
        title = result.title
        title = title.strip().strip('"')
        return title if title else "Chat Session"
    except Exception as e:
        st.error(f"Title generation failed: {e}")
        return "Chat Session"

class StreamlitCallbackHandler(BaseCallbackHandler):
    def __init__(self, status):
        self.status = status
        self.action = None

    def on_agent_action(self, action, **kwargs):
        self.action = action
        tool_input = action.tool_input
        query = "details"
        tool_name = action.tool

        if isinstance(tool_input, dict):
            query = tool_input.get("query", tool_input.get("input", str(tool_input)))
        elif isinstance(tool_input, str):
             query = tool_input

        if tool_name == "document_search":
             self.status.update(label=f"Searching document for: `{query}`")
        elif tool_name == "tavily_search_results_json":
             self.status.update(label=f"Searching the web for: `{query}`")
        elif tool_name == "qa_generation":
             num_pairs = tool_input.get('number', 'some') if isinstance(tool_input, dict) else 'some'
             self.status.update(label=f"Generating {num_pairs} Q&A pairs...")
        else:
             self.status.update(label=f"Using tool `{tool_name}`...")


    def on_tool_start(self, serialized, input_str, **kwargs):
        tool_name = serialized.get("name", "tool")
        query = "details"
        if self.action:
            tool_input = self.action.tool_input
            if isinstance(tool_input, dict):
                query = tool_input.get("query", tool_input.get("input", str(tool_input)))
            elif isinstance(tool_input, str):
                query = tool_input

        if tool_name == "document_search":
            self.status.update(label=f"Processing document search for: `{query}`")
        elif tool_name == "tavily_search_results_json":
            self.status.update(label=f"Processing web search for: `{query}`")
        elif tool_name == "qa_generation":
            self.status.update(label=f"Processing Q&A generation...")
        else:
            self.status.update(label=f"Processing with `{tool_name}`...")


    def on_tool_end(self, output, **kwargs):
        tool_name = self.action.tool if self.action else "tool"
        if tool_name == "document_search":
            self.status.update(label="Aggregating document search results")
        elif tool_name == "tavily_search_results_json":
            self.status.update(label="Aggregating web search results")
        elif tool_name == "qa_generation":
            self.status.update(label="Finalizing Q&A generation")
        else:
            self.status.update(label=f"Finished using `{tool_name}`")

def initialize_session_state():
    default_session_state = {
        "logged_in": False,
        "email": "",
        "username": "",
        "current_session_id": None,
        "current_session_title": "",
        "needs_title": False,
        "session_selected": False,
        "processed_file_id": None,
        "file_docs": None,
        "tools": [get_search_tool()],
        "downloadable_csv": None
    }
    for key, default_value in default_session_state.items():
        if key not in st.session_state:
            st.session_state[key] = default_value

def activate_latest_chat():
    """Activates the most recent chat session if none is selected."""
    if not st.session_state.session_selected:
        try:
            sessions = list(get_chat_sessions(user_id=st.session_state.email))
            if sessions:
                # Get the last session (most recent)
                latest_session_name = sessions[0][1]  # [1] gets the unique_name
                latest_session_id = sessions[0][0]    # [0] gets the session_id

                # Update session state
                st.session_state.current_session_id = latest_session_id
                st.session_state.current_session_title = latest_session_name
                st.session_state.session_selected = True
                st.session_state.needs_title = False
                st.session_state.processed_file_id = None
                st.session_state.file_docs = None
                st.session_state.tools = [get_search_tool()]

                # Update URL parameter
                st.query_params.s = latest_session_name
                return True
        except Exception as e:
            st.error(f"Failed to activate last chat: {e}")
    return False

initialize_session_state()
inject_custom_css()

if not st.user.is_logged_in:
    col1, col_main, col3 = st.columns([1, 5, 1])

    with col_main:
        st.header(f"Agent Lambda: AI Assistant", anchor=False)

        with st.container(border=True):
            st.markdown("<br>", unsafe_allow_html=True)
            google_logo_url = "https://upload.wikimedia.org/wikipedia/commons/c/c1/Google_%22G%22_logo.svg"
            st.markdown(
                f"""
                <div style="display: flex; justify-content: center; margin-top: 20px; margin-bottom: 20px;">
                    <img src="{google_logo_url}" width="25" height="25" style="margin-right: 10px; vertical-align: middle;">
                </div>
                """,
                unsafe_allow_html=True,
            )

            btn_col1, btn_col2, btn_col3 = st.columns([1, 2, 1])
            with btn_col2:
                login_pressed = st.button("Continue with Google", use_container_width=True)
                if login_pressed:
                    try:
                        st.login()
                    except Exception as e:
                        st.error(f"Login setup error: {e}. Ensure Google OAuth is configured.")

            st.markdown("<br>", unsafe_allow_html=True)

    st.stop()
else:
    if not st.session_state.logged_in or st.session_state.email != st.user.email:
        st.session_state.logged_in = True
        st.session_state.email = st.user.email
        st.session_state.username = st.user.name
        st.session_state.current_session_id = None
        st.session_state.current_session_title = ""
        st.session_state.needs_title = False
        st.session_state.session_selected = False
        st.session_state.processed_file_id = None
        st.session_state.file_docs = None
        st.session_state.tools = [get_search_tool()]
        st.session_state.downloadable_csv = None
        # Add this line to activate last chat after login
        activate_latest_chat()


with st.sidebar:
    st.write(f"Welcome, {st.session_state.username}")

    if st.button("➕ New Chat"):
        try:
            new_session_id, unique_name = create_chat_session(
                user_id=st.session_state.email,
                session_name="New Chat"
            )
            # Update session state
            st.session_state.current_session_id = new_session_id
            st.session_state.current_session_title = unique_name
            st.session_state.needs_title = True
            st.session_state.session_selected = True
            st.session_state.processed_file_id = None
            st.session_state.file_docs = None
            st.session_state.tools = [get_search_tool()]
            st.session_state.downloadable_csv = None

            # Update URL parameter immediately
            st.query_params.s = unique_name

            # Remove success message to avoid delay
            st.rerun()
        except Exception as e:
            st.error(f"Failed to create new chat: {e}")


    st.divider()
    query_params = st.query_params
    session_from_url = st.query_params.get("s", None)

    try:
        sessions_list = list(get_chat_sessions(user_id=st.session_state.email))
    except Exception as e:
        st.error(f"Failed to load chat sessions: {e}")
        sessions_list = []

    if sessions_list:
        session_options = {unique_name: sid for sid, unique_name in sessions_list}
        session_unique_names = [unique_name for sid, unique_name in sessions_list]

        st.markdown("## Chats")

        # First check URL parameter
        if session_from_url and session_from_url in session_options:
            if session_from_url != st.session_state.current_session_title:
                st.session_state.current_session_id = session_options[session_from_url]
                st.session_state.current_session_title = session_from_url
                st.session_state.needs_title = False
                st.session_state.session_selected = True
                st.session_state.processed_file_id = None
                st.session_state.file_docs = None
                st.session_state.tools = [get_search_tool()]
                st.session_state.downloadable_csv = None
        # If no URL parameter and no active session, activate the last chat
        elif not st.session_state.session_selected:
            activate_latest_chat()

        for session_name in session_unique_names:
            base_title = get_base_title(session_name)
            is_active = session_name == st.session_state.current_session_title

            if st.button(
                base_title,
                key=f"chat_{session_name}",
                use_container_width=True,
                type="secondary" if is_active else "tertiary",
            ):
                # Set query param with new session name
                st.query_params.s = session_name
                st.rerun()

    else:
        try:
            new_session_id, unique_name = create_chat_session(
                user_id=st.session_state.email,
                session_name="New Chat"
            )
            # Update session state
            st.session_state.current_session_id = new_session_id
            st.session_state.current_session_title = unique_name
            st.session_state.needs_title = True
            st.session_state.session_selected = True
            st.session_state.processed_file_id = None
            st.session_state.file_docs = None
            st.session_state.tools = [get_search_tool()]
            st.session_state.downloadable_csv = None

            # Update URL parameter immediately
            st.query_params.s = unique_name

            # Remove success message to avoid delay
            st.rerun()
        except Exception as e:
            st.error(f"Failed to create new chat: {e}")



if st.session_state.logged_in and st.session_state.session_selected and st.session_state.current_session_id:
    current_title_str = str(st.session_state.get("current_session_title", ""))
    display_title = get_base_title(current_title_str)
    st.subheader(f"Chat: {display_title}", anchor=False)

    try:
        session_id_obj = st.session_state.current_session_id
        if not isinstance(session_id_obj, ObjectId):
             session_id_obj = ObjectId(session_id_obj)

        messages = prepare_chat_history(
            session_id=session_id_obj,
            chat_history_limit=50
        )
    except Exception as e:
        st.error(f"Failed to load chat history: {e}")
        messages = []

    for message in messages:
        role = getattr(message, 'kind', getattr(message, 'type', 'ai' if hasattr(message, 'ai') else 'human'))
        display_role = "user" if role in ["human", "user"] else "assistant"
        with st.chat_message(display_role):
             st.markdown(str(message.content))

    if st.session_state.get('downloadable_csv'):
        csv_data = st.session_state['downloadable_csv']
        st.download_button(
            label="Download QA Pairs as CSV",
            data=csv_data["content"],
            file_name=csv_data["filename"],
            mime=csv_data["mime"],
            key=f"download_{csv_data['filename']}"
        )
        st.session_state.downloadable_csv = None


    prompt = st.chat_input(
        "Ask your questions or attach files...",
        accept_file=True,
        file_type=["pdf", "docx", "doc", "jpg", "jpeg", "png", "gif"]
    )

    if prompt:
        session_id_to_use = st.session_state.current_session_id
        if not isinstance(session_id_to_use, ObjectId):
            try:
                session_id_to_use = ObjectId(session_id_to_use)
            except Exception as e:
                st.error(f"Internal Error: Invalid session ID format. {e}")
                st.stop()

        # Handle text input
        prompt_text = prompt.text if prompt.text else "Uploaded files"
        if prompt.text:
            st.chat_message("user").markdown(prompt.text)

            # Title generation logic
            if st.session_state.needs_title:
                try:
                    with st.spinner("Generating title"):
                        base_title = generate_title_llm(prompt.text)
                    new_unique_title = update_session_name(session_id_to_use, base_title)
                    if new_unique_title:
                        st.session_state.current_session_title = new_unique_title
                        st.session_state.needs_title = False
                    else:
                        st.warning("Failed to update title, keeping old title.")
                        st.session_state.needs_title = False
                except Exception as title_e:
                    st.error(f"Failed to generate title: {title_e}")
                    st.session_state.needs_title = False

        # Handle file uploads
        if prompt.files:
            with st.status("Processing uploaded files...", expanded=True) as status:
                for uploaded_file in prompt.files:
                    status.update(label=f"Processing {uploaded_file.name}...")

                    try:
                        loader = LambdaStreamlitLoader(uploaded_file)
                        docs = list(loader.lazy_load())

                        if docs:
                            # Display images in chat
                            if any(doc.metadata.get("type") == "image" for doc in docs):
                                st.chat_message("user").markdown(
                                    f"Uploaded image: {uploaded_file.name}\n\n" +
                                    "\n".join(doc.page_content for doc in docs if doc.metadata.get("type") == "image")
                                )
                            else:
                                st.chat_message("user").markdown(f"Uploaded file: {uploaded_file.name}")

                            # Update document context
                            if not st.session_state.file_docs:
                                st.session_state.file_docs = docs
                            else:
                                st.session_state.file_docs.extend(docs)

                            st.session_state.processed_file_id = uploaded_file.file_id

                            # Update tools
                            current_tools = [get_search_tool()]
                            rag_tool = create_rag_tool(st.session_state.file_docs)
                            if rag_tool:
                                current_tools.append(rag_tool)
                            current_tools.append(qa_generation)
                            st.session_state.tools = current_tools

                            status.update(label=f"Successfully processed {uploaded_file.name}", state="complete")
                        else:
                            status.update(label=f"No content extracted from {uploaded_file.name}", state="error")

                    except Exception as e:
                        status.update(label=f"Error processing {uploaded_file.name}: {e}", state="error")

        # Save message to database
        try:
            add_message_to_session(
                session_id=session_id_to_use,
                content=prompt_text,
                kind="user"
            )
        except Exception as e:
            st.error(f"Failed to save message: {e}")

        # Process message with agent
        try:
            messages = prepare_chat_history(
                session_id=session_id_to_use,
                chat_history_limit=50
            )
        except Exception as e:
            st.error(f"Failed to reload chat history: {e}")
            messages = []

        agent_input = {"input": prompt_text, "chat_history": messages}
        current_tools = st.session_state.get("tools", [get_search_tool()])

        try:
            agent_executor = get_agent_executor(tools=current_tools)
        except Exception as agent_init_e:
            st.error(f"Failed to initialize the AI agent: {agent_init_e}")
            st.stop()

        with st.status("Processing your request...", expanded=False) as status:
            try:
                callback_handler = StreamlitCallbackHandler(status)
                response = agent_executor.invoke(
                    agent_input,
                    config={"callbacks": [callback_handler]}
                )

                output = response.get("output", "Sorry, I couldn't process that.")
                status.update(label="Done!", state="complete", expanded=True)

                with st.chat_message("assistant"):
                    st.markdown(str(output))

                try:
                    add_message_to_session(
                        session_id=session_id_to_use,
                        content=str(output),
                        kind="ai"
                    )
                except Exception as e:
                    st.error(f"Failed to save AI response: {e}")

                st.rerun()

            except Exception as e:
                error_message = f"An error occurred while processing your request: {e}. Please try again."
                if "rate limit" in str(e).lower() or "429" in str(e):
                    error_message = "Apologies, the system is experiencing high load (rate limit exceeded). Please try again in a few moments."
                elif "API key not valid" in str(e):
                    error_message = "Configuration error: An API key is invalid. Please contact support."

                status.update(label=f"Error: Processing failed.", state="error", expanded=True)
                st.chat_message("assistant").error(error_message)
